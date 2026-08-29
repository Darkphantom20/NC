from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_full_report_workflow():
    print("--- JRMSU Narrative Report Interactive Generator ---")
    print("Guide: Sections 1 to 6 (Acknowledgement, Introduction, Company Analysis, Tasks & Duties, Case Analysis, & Reflections).\n")
    
    # --- SECTION 1: ACKNOWLEDGEMENT ---
    print("=== SECTION 1: ACKNOWLEDGEMENT ===")
    student_name = input("Enter your full name [e.g., Ian P. Padilla]: ") or "Ian P. Padilla"
    degree_program = input("Enter your degree program [e.g., Bachelor of Science in Information System]: ") or "Bachelor of Science in Information System"
    
    print("\nEnter your acknowledgement paragraphs one by one.")
    print("Type 'done' or press Enter on an empty line when you are finished with Acknowledgement.\n")
    
    ack_paragraphs = []
    count = 1
    while True:
        text = input(f"Acknowledgement Paragraph {count} (or type 'done'): ")
        if text.strip().lower() == 'done' or text.strip() == '':
            if len(ack_paragraphs) == 0 and text.strip() == '':
                ack_paragraphs.append("With deepest gratitude and appreciation, I humbly extend my sincere thanks to all who contributed to my OJT.")
                print("Using default acknowledgement paragraph.")
            break
        ack_paragraphs.append(text)
        count += 1

    # --- SECTION 2: INTRODUCTION ---
    print("\n=== SECTION 2: INTRODUCTION ===")
    bg_org = input("1. Background of the Organization:\n> ") or \
             "The Management Information Systems Office (MISO) serves as the core technological backbone, responsible for managing, maintaining, and securing the digital infrastructure, information systems, and network communications of the institution."
             
    vision = input("2. Vision:\n> ") or \
             "To be a premier provider of innovative, reliable, and secure technological solutions and digital services."
             
    mission = input("3. Mission:\n> ") or \
              "To empower the organization through efficient IT infrastructure, responsive technical support, and robust systems development."
              
    print("\n4. Objectives (Enter items one by one. Type 'done' when finished):")
    objectives_list = []
    obj_count = 1
    while True:
        obj_text = input(f"Objective {obj_count} (or type 'done'): ")
        if obj_text.strip().lower() == 'done' or (obj_count == 1 and obj_text.strip() == ''):
            if len(objectives_list) == 0:
                objectives_list.append("To streamline administrative processes and ensure system reliability.")
            break
        objectives_list.append(obj_text)
        obj_count += 1
         
    print("\n5. Core Values (Enter items one by one. Type 'done' when finished):")
    core_values_list = []
    cv_count = 1
    while True:
        cv_text = input(f"Core Value {cv_count} (or type 'done'): ")
        if cv_text.strip().lower() == 'done' or (cv_count == 1 and cv_text.strip() == ''):
            if len(core_values_list) == 0:
                core_values_list.append("Integrity, Excellence, and Commitment to Public Service.")
            break
        core_values_list.append(cv_text)
        cv_count += 1
                 
    print("\n6. Products and Services Offered (Enter items one by one. Type 'done' when finished):")
    services_list = []
    serv_count = 1
    while True:
        serv_text = input(f"Service item {serv_count} (or type 'done'): ")
        if serv_text.strip().lower() == 'done' or (serv_count == 1 and serv_text.strip() == ''):
            if len(services_list) == 0:
                services_list.append("Provide reliable technical support and ICT services.")
            break
        services_list.append(serv_text)
        serv_count += 1

    # --- SECTION 3: ORGANIZATION / COMPANY ANALYSIS ---
    print("\n=== SECTION 3: ORGANIZATION / COMPANY ANALYSIS ===")
    
    print("\n1. Strengths (Internal factors) - Enter paragraphs one by one. Type 'done' when finished:")
    strengths_list = []
    s_count = 1
    while True:
        s_text = input(f"Strength Paragraph {s_count} (or type 'done'): ")
        if s_text.strip().lower() == 'done' or (s_count == 1 and s_text.strip() == ''):
            if len(strengths_list) == 0:
                strengths_list.append("The organization possesses highly skilled and dedicated technical personnel capable of handling complex network and software demands.")
            break
        strengths_list.append(s_text)
        s_count += 1

    print("\n2. Weaknesses (Internal factors) - Enter paragraphs one by one. Type 'done' when finished:")
    weaknesses_list = []
    w_count = 1
    while True:
        w_text = input(f"Weakness Paragraph {w_count} (or type 'done'): ")
        if w_text.strip().lower() == 'done' or (w_count == 1 and w_text.strip() == ''):
            if len(weaknesses_list) == 0:
                weaknesses_list.append("Internal operations occasionally face limitations due to aging hardware upgrades and constrained resource allocations.")
            break
        weaknesses_list.append(w_text)
        w_count += 1

    print("\n3. Opportunities (External factors) - Enter paragraphs one by one. Type 'done' when finished:")
    opportunities_list = []
    o_count = 1
    while True:
        o_text = input(f"Opportunity Paragraph {o_count} (or type 'done'): ")
        if o_text.strip().lower() == 'done' or (o_count == 1 and o_text.strip() == ''):
            if len(opportunities_list) == 0:
                opportunities_list.append("There are significant external prospects for adopting advanced cloud solutions and process automation technologies.")
            break
        opportunities_list.append(o_text)
        o_count += 1

    print("\n4. Threats (External factors) - Enter paragraphs one by one. Type 'done' when finished:")
    threats_list = []
    t_count = 1
    while True:
        t_text = input(f"Threat Paragraph {t_count} (or type 'done'): ")
        if t_text.strip().lower() == 'done' or (t_count == 1 and t_text.strip() == ''):
            if len(threats_list) == 0:
                threats_list.append("External risks such as evolving cybersecurity vulnerabilities and potential data breaches remain constant challenges.")
            break
        threats_list.append(t_text)
        t_count += 1

    print("\n5. Recommendations for Improvement - Enter paragraphs one by one. Type 'done' when finished:")
    recommendations_list = []
    r_count = 1
    while True:
        r_text = input(f"Recommendation Paragraph {r_count} (or type 'done'): ")
        if r_text.strip().lower() == 'done' or (r_count == 1 and r_text.strip() == ''):
            if len(recommendations_list) == 0:
                recommendations_list.append("It is highly recommended to upgrade legacy infrastructure and conduct continuous staff training programs.")
            break
        recommendations_list.append(r_text)
        r_count += 1

    # --- SECTION 4: TASKS AND DUTIES ---
    print("\n=== SECTION 4: TASKS AND DUTIES ===")
    
    print("\n1. Assigned Tasks and Responsibilities - Enter paragraphs one by one. Type 'done' when finished:")
    tasks_list = []
    tk_count = 1
    while True:
        tk_text = input(f"Task Paragraph {tk_count} (or type 'done'): ")
        if tk_text.strip().lower() == 'done' or (tk_count == 1 and tk_text.strip() == ''):
            if len(tasks_list) == 0:
                tasks_list.append("Assigned tasks included troubleshooting office hardware issues, maintaining network cables, and updating internal database records.")
            break
        tasks_list.append(tk_text)
        tk_count += 1

    print("\n2. Duties and Procedures Conformed - Enter paragraphs one by one. Type 'done' when finished:")
    procedures_list = []
    pr_count = 1
    while True:
        pr_text = input(f"Procedure Paragraph {pr_count} (or type 'done'): ")
        if pr_text.strip().lower() == 'done' or (pr_count == 1 and pr_text.strip() == ''):
            if len(procedures_list) == 0:
                procedures_list.append("Followed strict IT ticketing protocols, observed safety compliance guidelines during hardware repairs, and conformed to daily attendance logs.")
            break
        procedures_list.append(pr_text)
        pr_count += 1

    # --- SECTION 5: CASE ANALYSIS ---
    print("\n=== SECTION 5: CASE ANALYSIS ===")
    
    print("\n1. Issue / Problem 1 Details:")
    issue1_desc = input("Description of Problem 1:\n> ") or "Encountered unexpected network downtime during a critical system update, disrupting office communications."
    issue1_strat = input("Strategy/Action undertaken to solve Problem 1:\n> ") or "Conducted a line check, isolated the faulty switch port, and switched the main office routing temporarily to a secondary backup line."

    print("\n2. Issue / Problem 2 Details:")
    issue2_desc = input("Description of Problem 2:\n> ") or "Faced compatibility errors when deploying a legacy database script onto the updated server environment."
    issue2_strat = input("Strategy/Action undertaken to solve Problem 2:\n> ") or "Debugged SQL syntax constraints, updated driver packages, and successfully refactored connection strings."

    print("\n3. Lessons Learned - Enter paragraphs one by one. Type 'done' when finished:")
    lessons_list = []
    l_count = 1
    while True:
        l_text = input(f"Lesson Paragraph {l_count} (or type 'done'): ")
        if l_text.strip().lower() == 'done' or (l_count == 1 and l_text.strip() == ''):
            if len(lessons_list) == 0:
                lessons_list.append("These situations taught the vital importance of systematic troubleshooting, maintaining proper system backups, and remaining calm under pressure.")
            break
        lessons_list.append(l_text)
        l_count += 1

    # --- SECTION 6: REFLECTIONS ---
    print("\n=== SECTION 6: REFLECTIONS ===")
    
    print("\n1. Self-Evaluation from the Learning Process Experienced - Enter paragraphs one by one. Type 'done' when finished:")
    self_eval_list = []
    se_count = 1
    while True:
        se_text = input(f"Self-Evaluation Paragraph {se_count} (or type 'done'): ")
        if se_text.strip().lower() == 'done' or (se_count == 1 and se_text.strip() == ''):
            if len(self_eval_list) == 0:
                self_eval_list.append("The OJT journey served as a transformative learning process, pushing me to transition from theoretical classroom knowledge to practical, fast-paced technical execution.")
            break
        self_eval_list.append(se_text)
        se_count += 1

    print("\n2. Relevancy of the Organization with Programme of Study & Expected Goals - Enter paragraphs one by one. Type 'done' when finished:")
    relevancy_list = []
    rv_count = 1
    while True:
        rv_text = input(f"Relevancy Paragraph {rv_count} (or type 'done'): ")
        if rv_text.strip().lower() == 'done' or (rv_count == 1 and rv_text.strip() == ''):
            if len(relevancy_list) == 0:
                relevancy_list.append("The host organization directly aligns with my degree program, allowing me to fulfill my expected professional goals of mastering enterprise systems administration and IT support workflows.")
            break
        relevancy_list.append(rv_text)
        rv_count += 1

    # --- DOCUMENT GENERATION ---
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.2)
    
    header = section.header
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p_head.add_run().add_picture('Picture1.png', width=Inches(7.0))
    except Exception:
        print("\nNote: Picture1.png not found, header image skipped.")
        
    # --- BUILD ACKNOWLEDGEMENT ---
    p_ack_title = doc.add_paragraph()
    p_ack_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ack_title.paragraph_format.space_before = Pt(18)
    p_ack_title.paragraph_format.space_after = Pt(12)
    r_ack_title = p_ack_title.add_run("ACKNOWLEDGEMENT")
    r_ack_title.bold = True
    r_ack_title.font.name = 'Times New Roman'
    r_ack_title.font.size = Pt(14)
    
    for text in ack_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
    doc.add_paragraph()
    
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.15
    
    r_name = p_sig.add_run(f"{student_name.upper()}\n")
    r_name.bold = True
    r_name.font.name = 'Times New Roman'
    r_name.font.size = Pt(11)
    
    r_prog = p_sig.add_run(f"{degree_program}\n")
    r_prog.font.name = 'Times New Roman'
    r_prog.font.size = Pt(11)
    
    r_univ = p_sig.add_run("Jose Rizal Memorial State University")
    r_univ.font.name = 'Times New Roman'
    r_univ.font.size = Pt(11)
    
    # --- BUILD INTRODUCTION ---
    doc.add_page_break()
    p_intro_title = doc.add_paragraph()
    p_intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_intro_title.paragraph_format.space_before = Pt(18)
    p_intro_title.paragraph_format.space_after = Pt(14)
    r_intro_title = p_intro_title.add_run("2. INTRODUCTION")
    r_intro_title.bold = True
    r_intro_title.font.name = 'Times New Roman'
    r_intro_title.font.size = Pt(14)
    
    intro_blocks = [
        ("Background of the Organization", [bg_org]),
        ("Vision", [vision]),
        ("Mission", [mission]),
        ("Objectives", objectives_list),
        ("Core Values", core_values_list),
        ("Products and Services Offered", services_list)
    ]
    
    for subtitle, lines in intro_blocks:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(10)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(12)
        
        for line in lines:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(4)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if len(lines) > 1 and subtitle in ["Objectives", "Core Values", "Products and Services Offered"]:
                p_body.style = 'List Bullet'
                
            r_body = p_body.add_run(line)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(11)

    # --- BUILD COMPANY ANALYSIS ---
    doc.add_page_break()
    p_ca_title = doc.add_paragraph()
    p_ca_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ca_title.paragraph_format.space_before = Pt(18)
    p_ca_title.paragraph_format.space_after = Pt(14)
    r_ca_title = p_ca_title.add_run("3. ORGANIZATION / COMPANY ANALYSIS")
    r_ca_title.bold = True
    r_ca_title.font.name = 'Times New Roman'
    r_ca_title.font.size = Pt(14)
    
    analysis_blocks = [
        ("Strengths of the Organization (Internal factors)", strengths_list),
        ("Weaknesses of the Organization (Internal factors)", weaknesses_list),
        ("Opportunities of the Organization (External factors)", opportunities_list),
        ("Threats of the Organization (External factors)", threats_list),
        ("Recommendations for Improvement", recommendations_list)
    ]
    
    for subtitle, paragraphs in analysis_blocks:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(10)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(12)
        
        for para_text in paragraphs:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(6)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            r_body = p_body.add_run(para_text)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(11)

    # --- BUILD TASKS AND DUTIES ---
    doc.add_page_break()
    p_td_title = doc.add_paragraph()
    p_td_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_td_title.paragraph_format.space_before = Pt(18)
    p_td_title.paragraph_format.space_after = Pt(14)
    r_td_title = p_td_title.add_run("4. TASKS AND DUTIES")
    r_td_title.bold = True
    r_td_title.font.name = 'Times New Roman'
    r_td_title.font.size = Pt(14)
    
    task_blocks = [
        ("Assigned Tasks and Responsibilities", tasks_list),
        ("Duties and Procedures Conformed", procedures_list)
    ]
    
    for subtitle, paragraphs in task_blocks:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(10)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(12)
        
        for para_text in paragraphs:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(6)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            r_body = p_body.add_run(para_text)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(11)

    # --- BUILD CASE ANALYSIS ---
    doc.add_page_break()
    p_cs_title = doc.add_paragraph()
    p_cs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cs_title.paragraph_format.space_before = Pt(18)
    p_cs_title.paragraph_format.space_after = Pt(14)
    r_cs_title = p_cs_title.add_run("5. CASE ANALYSIS")
    r_cs_title.bold = True
    r_cs_title.font.name = 'Times New Roman'
    r_cs_title.font.size = Pt(14)
    
    case_blocks = [
        ("Issue / Problem 1", [f"Description: {issue1_desc}", f"Strategy/Action Undertaken: {issue1_strat}"]),
        ("Issue / Problem 2", [f"Description: {issue2_desc}", f"Strategy/Action Undertaken: {issue2_strat}"]),
        ("Lessons Learned from the Situations", lessons_list)
    ]
    
    for subtitle, paragraphs in case_blocks:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(10)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(12)
        
        for para_text in paragraphs:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(6)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            r_body = p_body.add_run(para_text)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(11)

    # --- BUILD REFLECTIONS ---
    doc.add_page_break()
    p_ref_title = doc.add_paragraph()
    p_ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ref_title.paragraph_format.space_before = Pt(18)
    p_ref_title.paragraph_format.space_after = Pt(14)
    r_ref_title = p_ref_title.add_run("6. REFLECTIONS")
    r_ref_title.bold = True
    r_ref_title.font.name = 'Times New Roman'
    r_ref_title.font.size = Pt(14)
    
    reflection_blocks = [
        ("Self-Evaluation from the Learning Process Experienced", self_eval_list),
        ("Relevancy of the Organization with Your Programme of Study and Expected Goals", relevancy_list)
    ]
    
    for subtitle, paragraphs in reflection_blocks:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(10)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(12)
        
        for para_text in paragraphs:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(6)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            r_body = p_body.add_run(para_text)
            r_body.font.name = 'Times New Roman'
            r_body.font.size = Pt(11)
        
    filename = "JRMSU_Narrative_Report_Sections_1_to_6.docx"
    doc.save(filename)
    print(f"\nSuccessfully generated and saved your document through Section 6 to '{filename}'.")

if __name__ == "__main__":
    create_full_report_workflow()